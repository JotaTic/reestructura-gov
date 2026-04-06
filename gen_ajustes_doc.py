#!/usr/bin/env python3
"""Generate Ajustes.docx - complete improvement prompt for ReEstructura.Gov."""
import json
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('ReEstructura.Gov — Prompt de Mejora Integral', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Documento de análisis, nomenclatura completa y mejoras requeridas para la plataforma '
    'de reestructuración de entidades públicas colombianas.\n'
    'Fecha: Abril 2026 | Dominio: reestructuracion.corpofuturo.org'
)

# ============================================================
# PART 1: NOMENCLATURA COMPLETA
# ============================================================
doc.add_heading('PARTE 1: NOMENCLATURA COMPLETA DE EMPLEOS', level=1)

doc.add_paragraph(
    'Catálogo completo extraído del Decreto 785 de 2005 (territorial) y '
    'Decreto 2489 de 2006 (nacional), complementado con el Concepto 149331 de 2022 del DAFP '
    'que aclara que las entidades territoriales SOLO pueden usar las denominaciones del '
    'Decreto 785/2005 y no pueden crear denominaciones nuevas fuera de esta tabla.'
)

# Load fixture
with open('plataforma/backend/apps/nomenclatura/fixtures/decreto_785_2005.json', 'r', encoding='utf-8') as f:
    fixture = json.load(f)

# Group by scope and level
from collections import defaultdict
grouped = defaultdict(lambda: defaultdict(list))
for rec in fixture:
    f = rec['fields']
    grouped[f['scope']][f['level']].append((f['code'], f['denomination']))

level_order = ['DIRECTIVO', 'ASESOR', 'PROFESIONAL', 'TECNICO', 'ASISTENCIAL']
scope_titles = {
    'TERRITORIAL': 'ORDEN TERRITORIAL — Decreto 785 de 2005',
    'NACIONAL': 'ORDEN NACIONAL — Decreto 2489 de 2006',
}

for scope in ['TERRITORIAL', 'NACIONAL']:
    doc.add_heading(scope_titles[scope], level=2)

    for level in level_order:
        items = grouped[scope].get(level, [])
        if not items:
            continue

        doc.add_heading(f'Nivel {level.title()} ({len(items)} denominaciones)', level=3)

        # Table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        hdr = table.rows[0].cells
        hdr[0].text = 'Código'
        hdr[1].text = 'Denominación del Empleo'
        for cell in hdr:
            for p in cell.paragraphs:
                p.style.font.bold = True

        # Sort by code
        items.sort(key=lambda x: x[0])

        for code, denom in items:
            row = table.add_row().cells
            row[0].text = code
            row[1].text = denom

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(13)

        doc.add_paragraph('')  # spacing

# Summary table
doc.add_heading('Resumen de Nomenclatura', level=2)
summary_table = doc.add_table(rows=1, cols=3)
summary_table.style = 'Light Grid Accent 1'
hdr = summary_table.rows[0].cells
hdr[0].text = 'Ámbito'
hdr[1].text = 'Nivel'
hdr[2].text = 'Cantidad'

for scope in ['TERRITORIAL', 'NACIONAL']:
    for level in level_order:
        items = grouped[scope].get(level, [])
        row = summary_table.add_row().cells
        row[0].text = scope
        row[1].text = level
        row[2].text = str(len(items))

row = summary_table.add_row().cells
row[0].text = 'TOTAL'
row[1].text = ''
row[2].text = str(len(fixture))

doc.add_paragraph('')

# ============================================================
# PART 2: ANALYSIS PER MODULE
# ============================================================
doc.add_heading('PARTE 2: ANÁLISIS POR MÓDULO', level=1)

modules = [
    ('Estructura Orgánica',
     'Solo lectura. Muestra árbol de dependencias. Los cargos se crean en Planta de Personal.',
     ['Editor visual de organigrama con drag & drop',
      'Click en dependencia para ver cargos asignados',
      'Crear/editar dependencias inline con botón "Agregar subdependencia"',
      'Vista comparativa: organigrama actual vs propuesto lado a lado',
      'Indicador visual: verde (con cargos), amarillo (sin cargos), rojo (sin funciones)']),

    ('Hojas de Vida (Talento)',
     'Creación individual por formulario + importación masiva SIGEP (Excel). No permite adjuntar Word/PDF.',
     ['Adjuntar documento PDF/Word como hoja de vida (FileField en Employee)',
      'Guardar en media/hojas_de_vida/{entity_id}/{employee_id}/',
      'Vista individual completa: datos + educación + experiencia + capacitación + evaluaciones',
      'Editar empleado existente (actualmente solo se puede ver)',
      'Exportar hoja de vida individual a DOCX']),

    ('Manual de Funciones Vigente',
     'Permite crear manual + importar DOCX que parsea identificación, propósito, funciones y requisitos.',
     ['Mejorar parser DOCX para detectar más formatos de manual',
      'Permitir editar roles/funciones importados',
      'Vincular roles del manual vigente con cargos de la planta actual',
      'Soporte para PDF además de DOCX']),

    ('Manual Específico de Funciones (Propuesto)',
     'Auto-generado desde matriz de cargas. Solo lectura, no editable.',
     ['Permitir editar el manual generado antes de exportar',
      'Ajustar propósito, funciones y requisitos por cargo',
      'Guardar versión editada sin perder la auto-generada',
      'Modelo ManualFuncionesOverride para customizaciones']),

    ('Procedimientos',
     'Importa DOCX con pasos. No distingue actual vs propuesto. No edita pasos post-importación.',
     ['Agregar campo kind (VIGENTE/PROPUESTO) al modelo Procedure',
      'Permitir editar pasos después de importar (reordenar, modificar, agregar/eliminar)',
      'Botón "Clonar como propuesto" para crear versión ajustada',
      'Vista comparativa vigente vs propuesto con diff visual',
      'Dropdown de procesos en lugar de ingresar ID manualmente']),

    ('Base Legal vs Mandatos Legales',
     'Base Legal = biblioteca de referencia (leyes, decretos, sentencias). '
     'Mandatos = obligaciones específicas por entidad extraídas de esas normas, vinculables a procesos.',
     ['Base Legal: permitir agregar normas personalizadas además de las seeded',
      'Mandatos: UI matricial para vincular mandatos con procesos (existe modelo MandateCompliance pero sin UI)',
      'Mostrar reporte de brechas en frontend: mandatos sin proceso, procesos sin mandato',
      'Porcentaje global de cobertura de mandatos']),

    ('Matriz de Cargas de Trabajo',
     'Formulario 1 Anexo 5 DAFP. Entrada manual actividad por actividad. Sin importación Excel.',
     ['Importación masiva desde Excel (columnas: proceso, actividad, frecuencia, tiempos)',
      'Cálculo automático de tiempo estándar TE = [(Tmin + 4×TU + Tmax) / 6] × 1.07',
      'Consolidado automático por nivel y departamento (Formulario 2)',
      'Poder importar matriz de cargas existente (no solo crear desde cero)',
      'Exportar a Excel con formato oficial DAFP']),

    ('Planta de Personal',
     'CRUD de planes (ACTUAL/PROPUESTO) + posiciones. Sin importación masiva.',
     ['Importar planta actual desde Excel (dependencia, nivel, denominación, código, grado, cantidad, salario)',
      'Asociar automáticamente a dependencia por nombre',
      'Aplicar escala salarial automáticamente al importar',
      'Comparación visual actual vs propuesta con deltas resaltados']),

    ('Elegibilidad',
     'Análisis bulk por nivel origen → destino. Evalúa educación + experiencia por Decreto 785/2005.',
     ['Análisis individual: seleccionar UN empleado y ver elegibilidad para TODOS los niveles',
      'Mostrar "hoja de ruta" para NO_ELEGIBLE (qué le falta y cuánto tiempo/estudio)',
      'Exportar resultado a Excel con todos los empleados',
      'Vincular con cargo específico (no solo nivel genérico)']),

    ('Simulador',
     'Clona plan → evalúa métricas (posiciones, costo, Ley 617, mandatos) → compara escenarios.',
     ['Editar posiciones del escenario inline (tabla editable)',
      'Recalcular métricas al guardar cambios',
      'Escenarios automáticos sugeridos: "mínimo legal", "optimizado Ley 617"',
      'Exportar comparación a XLSX/DOCX/PDF',
      'Gráficas de comparación (barras de costo, semáforos Ley 617)']),

    ('Gobierno del Proceso',
     'Máquina de estados: BORRADOR → DIAGNÓSTICO → ANÁLISIS → REVISIÓN JURÍDICA → '
     'REVISIÓN FINANCIERA → CONCEPTO DAFP → COMISIÓN PERSONAL → APROBADO → ACTO EXPEDIDO → '
     'IMPLEMENTADO → ARCHIVADO. Cada transición tiene precondiciones.',
     ['Dashboard visual del avance por fase con porcentaje de completitud',
      'Timeline horizontal mostrando duración en cada fase',
      'Alertas cuando una fase lleva demasiado tiempo',
      'Vincular evidencias (documentos) a cada transición']),

    ('Actos Administrativos',
     'Genera borradores desde plantillas: decreto estructura, planta, manual funciones, '
     'escala salarial, supresión, liquidación.',
     ['Preview en tiempo real del borrador con placeholders resueltos',
      'Flujo de revisión: BORRADOR → EN REVISIÓN → APROBADO → EXPEDIDO con comentarios',
      'Generar PDF con formato oficial (membrete, numeración)',
      'Vincular al workflow: al expedir acto, avanzar estado automáticamente',
      'Gestión de plantillas (crear/editar) desde el frontend']),

    ('Nomenclatura',
     f'Catálogo oficial ahora con 290 denominaciones (antes 49). Solo lectura, sin importación.',
     ['Importación desde Excel para agregar nomenclatura personalizada',
      'Marcar denominaciones "custom" vs "oficial"',
      'Filtrar por nivel con tabs: Directivo | Asesor | Profesional | Técnico | Asistencial',
      'Incluir grados y rangos salariales por denominación']),

    ('Comisión de Personal',
     'CRUD de comisiones y reuniones. Miembros almacenados como JSON sin UI de gestión.',
     ['CRUD de miembros con nombre, cargo, tipo (representante empleados/entidad), período',
      'Adjuntar PDF del acta de reunión firmada',
      'Vincular reuniones como evidencia para paso COMISION_PERSONAL_INFORMADA del workflow']),

    ('Análisis Financiero + MFMP',
     'Tabla editable de años fiscales con indicadores Ley 617/Ley 358. MFMP con proyección a N años.',
     ['Gráficas de proyección: líneas de ingresos, gastos, deuda por año',
      'Semáforos por año proyectado para Ley 617 y Ley 358',
      'Preview de datos antes de confirmar importación FUT',
      'Ficha de impacto fiscal con formato oficial']),

    ('Retén Social',
     'Registro de protegidos con sincronización automática desde empleados.',
     ['Dashboard de protegidos: gráfica por tipo de protección',
      'Alertas de vencimiento de protecciones (pre-pensión, embarazo, lactancia)',
      'Exportar con detalle legal por tipo de protección']),
]

for name, current, improvements in modules:
    doc.add_heading(name, level=2)
    doc.add_paragraph(f'Estado actual: {current}')
    doc.add_heading('Mejoras requeridas:', level=4)
    for imp in improvements:
        doc.add_paragraph(imp, style='List Bullet')

# ============================================================
# PART 3: CROSS-CUTTING IMPROVEMENTS
# ============================================================
doc.add_heading('PARTE 3: MEJORAS TRANSVERSALES', level=1)

transversal = [
    ('Importación masiva desde Excel',
     'Crear endpoint POST /api/<recurso>/importar-xlsx/ para: Planta, Matriz de Cargas, '
     'Nomenclatura, Escala Salarial. Patrón: aceptar .xlsx, validar headers, transacción atómica, '
     'retornar {created, updated, errors[], warnings[]}.'),
    ('Búsqueda global',
     'Barra de búsqueda en Topbar que busque en entidades, empleados, dependencias, cargos, mandatos.'),
    ('Auditoría visible',
     'En cada registro mostrar quién lo creó/modificó (ya existe AuditedModel, falta exponerlo en UI).'),
    ('Estudio técnico unificado',
     'Botón "Generar Estudio Técnico Completo" que compile TODOS los módulos en un solo DOCX/PDF '
     'con tabla de contenido: diagnóstico, marco legal, DOFA, procesos, cargas, planta, manual, '
     'financiero, retén social, concepto DAFP, actos.'),
    ('Notificaciones en tiempo real',
     'WebSocket o polling para cambios de estado, respuestas a consultas, vencimientos.'),
    ('Responsive mejorado',
     'Tablas con scroll horizontal o cards en mobile. Varias tablas no se ven bien en móvil.'),
]

for name, desc in transversal:
    doc.add_heading(name, level=2)
    doc.add_paragraph(desc)

# ============================================================
# PART 4: PRIORITIZATION
# ============================================================
doc.add_heading('PARTE 4: PRIORIZACIÓN', level=1)

sprints = [
    ('Sprint 7 — Importaciones (ALTA prioridad)',
     ['Importación Excel de Planta de Personal',
      'Importación Excel de Matriz de Cargas',
      'Importación/personalización de Nomenclatura',
      'Adjuntar documento HV a empleado',
      'Cargar fixture completo de nomenclatura (290 registros)']),
    ('Sprint 8 — Edición y comparación',
     ['Editor visual de estructura orgánica',
      'Edición post-generación del manual de funciones',
      'Procedimientos: actual vs propuesto + edición de pasos',
      'UI de vinculación mandatos ↔ procesos']),
    ('Sprint 9 — Simulación y reportes',
     ['Edición inline de escenarios + escenarios sugeridos',
      'Elegibilidad individual + exportar resultado',
      'Dashboard con gráficas',
      'Estudio técnico unificado (DOCX/PDF)']),
    ('Sprint 10 — Gobierno y actos',
     ['Preview de actos + flujo de revisión + PDF',
      'CRUD miembros comisión + adjuntar actas',
      'Dashboard retén social + alertas vencimiento']),
    ('Sprint 11 — Polish',
     ['Gráficas MFMP + preview FUT',
      'Búsqueda global',
      'Notificaciones en tiempo real',
      'Responsive mejorado']),
]

for name, items in sprints:
    doc.add_heading(name, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ============================================================
# PART 5: TECHNICAL NOTES
# ============================================================
doc.add_heading('PARTE 5: NOTAS TÉCNICAS', level=1)

notes = [
    'Stack: Django 5.0.6 + DRF 3.15 + PostgreSQL | Next.js 14 + Tailwind CSS',
    'Repo: github.com/JotaTic/reestructura-gov (branch: main)',
    'Deploy: ssh reestructura-droplet && bash deploy.sh',
    'Todos los imports deben usar openpyxl (ya en requirements.txt)',
    'Para PDF export agregar weasyprint o reportlab a requirements.txt',
    'Para organigrama interactivo usar reactflow (MIT license)',
    'Para gráficas de dashboard usar recharts (compatible con Next.js)',
    'Mantener patrón EntityScopedMixin / RestructuringScopedMixin',
    'Frontend: patrón api.get/api.post desde @/lib/api.ts',
    'Cada módulo con import necesita tests de parsing',
    'Nomenclatura actualizada: 290 registros (Dec 785/2005 + Dec 2489/2006)',
]

for note in notes:
    doc.add_paragraph(note, style='List Bullet')

# Save
output_path = 'Ajustes.docx'
doc.save(output_path)
print(f'Documento generado: {output_path}')
