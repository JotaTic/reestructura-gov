"""
Servicios para el módulo de Manual de Funciones Vigente (legacy).

- parse_manual_docx: parser heurístico con python-docx.
- parse_manual_pdf: placeholder (pypdf no instalado en Sprint 3).
- compare_current_vs_proposed: diff entre manual vigente y planta propuesta.
"""
from __future__ import annotations

import re
import unicodedata
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from .models import LegacyManual


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _normalize(text: str) -> str:
    """Convierte a ASCII-lower para comparaciones insensibles a tildes."""
    return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode().lower().strip()


def _parse_minutes(text: str) -> int:
    """Convierte strings como '15 min', '1 hora', '30' a minutos enteros."""
    t = text.lower().strip()
    if 'hora' in t:
        m = re.search(r'(\d+(?:\.\d+)?)', t)
        if m:
            return int(float(m.group(1)) * 60)
    m = re.search(r'(\d+)', t)
    return int(m.group(1)) if m else 0


# ---------------------------------------------------------------------------
# parse_manual_docx
# ---------------------------------------------------------------------------

def parse_manual_docx(file, manual: 'LegacyManual') -> dict:
    """
    Importador heurístico de manual de funciones en DOCX.

    Estrategia:
    1. Detecta "IDENTIFICACION DEL EMPLEO" → nuevo cargo.
    2. Extrae metadatos: nivel, código, grado, denominación.
    3. Detecta "PROPOSITO PRINCIPAL" → concatena párrafos.
    4. Detecta "FUNCIONES ESENCIALES" / "DESCRIPCION DE FUNCIONES" → crea LegacyManualFunction.
    5. Detecta "REQUISITOS" → separa ESTUDIOS y EXPERIENCIA.
    6. Fallback: si no hay cargos y hay tablas, intenta parsear la primera tabla.

    Persiste en DB y guarda reporte en manual.import_report.
    """
    from docx import Document as DocxDocument
    from .models import LegacyManualRole, LegacyManualFunction, RoleLevel

    doc = DocxDocument(file)
    warnings: list[str] = []

    # Collect all content blocks (paragraphs and table cell text) in document order
    # We use a state machine approach
    roles_created = 0
    functions_created = 0

    # Keywords for detection
    KW_IDENTIFICACION = 'identificacion del empleo'
    KW_PROPOSITO = 'proposito principal'
    KW_FUNCIONES_1 = 'funciones esenciales'
    KW_FUNCIONES_2 = 'descripcion de funciones'
    KW_REQUISITOS = 'requisitos'
    KW_ESTUDIOS = 'estudios'
    KW_EXPERIENCIA = 'experiencia'

    # State
    current_role_data: dict | None = None
    current_roles_data: list[dict] = []
    current_section: str = ''
    purpose_lines: list[str] = []
    function_lines: list[str] = []
    education_lines: list[str] = []
    experience_lines: list[str] = []
    in_requisitos = False
    in_estudios = False
    in_experiencia = False

    def flush_role():
        nonlocal current_role_data
        if current_role_data is not None:
            current_role_data['main_purpose'] = '\n'.join(purpose_lines).strip()
            current_role_data['function_lines'] = list(function_lines)
            current_role_data['min_education'] = '\n'.join(education_lines).strip()
            current_role_data['min_experience'] = '\n'.join(experience_lines).strip()
            current_roles_data.append(current_role_data)
        purpose_lines.clear()
        function_lines.clear()
        education_lines.clear()
        experience_lines.clear()

    def detect_section(norm_text: str) -> str | None:
        if KW_IDENTIFICACION in norm_text:
            return 'identificacion'
        if KW_PROPOSITO in norm_text:
            return 'proposito'
        if KW_FUNCIONES_1 in norm_text or KW_FUNCIONES_2 in norm_text:
            return 'funciones'
        if KW_REQUISITOS in norm_text:
            return 'requisitos'
        return None

    def parse_metadata_from_text(text: str) -> dict:
        """Intenta extraer metadatos del cargo de un bloque de texto."""
        meta = {'level': '', 'code': '', 'grade': '', 'denomination': ''}
        level_m = re.search(r'NIVEL[:\s]+(\w+)', text, re.IGNORECASE)
        code_m = re.search(r'C[OÓ]DIGO[:\s]+(\d+)', text, re.IGNORECASE)
        grade_m = re.search(r'GRADO[:\s]+(\d+)', text, re.IGNORECASE)
        denom_m = re.search(r'DENOMINACI[OÓ]N[^:]*:\s*(.+)', text, re.IGNORECASE)
        if level_m:
            meta['level'] = level_m.group(1).strip().upper()
        if code_m:
            meta['code'] = code_m.group(1).strip()
        if grade_m:
            meta['grade'] = grade_m.group(1).strip()
        if denom_m:
            meta['denomination'] = denom_m.group(1).strip()
        return meta

    def map_level(raw: str) -> str:
        mapping = {
            'DIRECTIVO': 'DIRECTIVO',
            'ASESOR': 'ASESOR',
            'PROFESIONAL': 'PROFESIONAL',
            'TECNICO': 'TECNICO',
            'ASISTENCIAL': 'ASISTENCIAL',
        }
        return mapping.get(raw.upper(), 'ASISTENCIAL')

    # --- Iterate document elements ---
    # We need to process paragraphs and tables in document order.
    # Use python-docx's element iteration.
    from docx.oxml.ns import qn

    for element in doc.element.body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # Paragraph
            text = ''.join(r.text or '' for r in element.iter(qn('w:r'))
                           if r.find(qn('w:t')) is not None)
            text = text.strip()
            if not text:
                continue
            norm = _normalize(text)

            section = detect_section(norm)
            if section == 'identificacion':
                flush_role()
                current_role_data = {'level': '', 'code': '', 'grade': '', 'denomination': '',
                                     'main_purpose': '', 'min_education': '', 'min_experience': '',
                                     'function_lines': []}
                current_section = 'identificacion'
                in_requisitos = False
                in_estudios = False
                in_experiencia = False
                continue
            elif section == 'proposito':
                current_section = 'proposito'
                in_requisitos = False
                continue
            elif section == 'funciones':
                current_section = 'funciones'
                in_requisitos = False
                continue
            elif section == 'requisitos':
                current_section = 'requisitos'
                in_requisitos = True
                in_estudios = False
                in_experiencia = False
                continue

            if current_role_data is None:
                continue

            if current_section == 'identificacion':
                meta = parse_metadata_from_text(text)
                for k, v in meta.items():
                    if v:
                        current_role_data[k] = v

            elif current_section == 'proposito':
                purpose_lines.append(text)

            elif current_section == 'funciones':
                # Any non-empty line becomes a function
                if text.strip():
                    function_lines.append(text)

            elif current_section == 'requisitos':
                norm_line = _normalize(text)
                if KW_ESTUDIOS in norm_line:
                    in_estudios = True
                    in_experiencia = False
                elif KW_EXPERIENCIA in norm_line:
                    in_estudios = False
                    in_experiencia = True
                elif in_estudios:
                    education_lines.append(text)
                elif in_experiencia:
                    experience_lines.append(text)

        elif tag == 'tbl':
            # Table — try to extract metadata or functions
            from docx.table import Table
            try:
                table = Table(element, doc)
                if current_role_data is not None and current_section == 'identificacion':
                    # Try to extract metadata from table rows
                    for row in table.rows:
                        row_text = ' | '.join(c.text.strip() for c in row.cells if c.text.strip())
                        if row_text:
                            meta = parse_metadata_from_text(row_text)
                            for k, v in meta.items():
                                if v and not current_role_data.get(k):
                                    current_role_data[k] = v
                elif current_role_data is not None and current_section == 'funciones':
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                function_lines.append(cell.text.strip())
            except Exception as e:
                warnings.append(f'Error procesando tabla: {e}')

    # Flush last role
    flush_role()

    # Fallback: if no roles found but there are tables, try first table as matrix
    if not current_roles_data and doc.tables:
        warnings.append('No se detectaron cargos por encabezados. Intentando fallback con primera tabla.')
        table = doc.tables[0]
        if len(table.columns) >= 3:
            headers = [_normalize(c.text) for c in table.rows[0].cells]
            col_map = {}
            for i, h in enumerate(headers):
                if 'nivel' in h:
                    col_map['level'] = i
                elif 'cod' in h:
                    col_map['code'] = i
                elif 'grad' in h:
                    col_map['grade'] = i
                elif 'denom' in h or 'empleo' in h:
                    col_map['denomination'] = i
                elif 'prop' in h or 'objeto' in h:
                    col_map['main_purpose'] = i

            for row in table.rows[1:]:
                cells = row.cells
                role_data = {
                    'level': cells[col_map.get('level', 0)].text.strip() if 'level' in col_map else '',
                    'code': cells[col_map.get('code', 1)].text.strip() if 'code' in col_map else '',
                    'grade': cells[col_map.get('grade', 2)].text.strip() if 'grade' in col_map else '',
                    'denomination': cells[col_map.get('denomination', 0)].text.strip() if 'denomination' in col_map else '',
                    'main_purpose': cells[col_map.get('main_purpose', 0)].text.strip() if 'main_purpose' in col_map else '',
                    'min_education': '',
                    'min_experience': '',
                    'function_lines': [],
                }
                if any(role_data[k] for k in ('code', 'grade', 'denomination')):
                    current_roles_data.append(role_data)

    # Persist to DB
    order = 0
    for rd in current_roles_data:
        level_raw = rd.get('level', 'ASISTENCIAL')
        level = map_level(level_raw)
        role = LegacyManualRole.objects.create(
            manual=manual,
            level=level,
            code=rd.get('code', '') or '',
            grade=rd.get('grade', '') or '',
            denomination=rd.get('denomination', '') or 'Sin denominación',
            main_purpose=rd.get('main_purpose', '') or '',
            dependencies_where_applies='',
            min_education=rd.get('min_education', '') or '',
            min_experience=rd.get('min_experience', '') or '',
            order=order,
        )
        roles_created += 1
        order += 1

        for fi, func_text in enumerate(rd.get('function_lines', [])):
            if func_text.strip():
                LegacyManualFunction.objects.create(
                    role=role,
                    order=fi,
                    description=func_text.strip(),
                    is_essential=True,
                )
                functions_created += 1

    report = {
        'roles_created': roles_created,
        'functions_created': functions_created,
        'warnings': warnings,
    }
    manual.import_report = report
    manual.save(update_fields=['import_report'])
    return report


# ---------------------------------------------------------------------------
# parse_manual_pdf
# ---------------------------------------------------------------------------

def parse_manual_pdf(file, manual: 'LegacyManual') -> dict:
    """
    Parser PDF — pendiente.

    Decisión Sprint 3: pypdf no se instala como dependencia obligatoria.
    El flujo principal soporta solo .docx en este sprint.
    Si en sprints posteriores se requiere PDF, instalar pypdf==4.3.1
    y extraer texto plano para reutilizar la lógica de parse_manual_docx.
    """
    return {
        'roles_created': 0,
        'functions_created': 0,
        'warnings': ['Parser PDF pendiente — pypdf no instalado en Sprint 3. Use .docx.'],
    }


# ---------------------------------------------------------------------------
# compare_current_vs_proposed
# ---------------------------------------------------------------------------

def compare_current_vs_proposed(entity, restructuring) -> dict:
    """
    Compara cargos del manual vigente más reciente de la entidad
    con los cargos del plan propuesto más reciente de la restructuring.

    Match por (code, grade) si no vacíos; fallback por denomination normalizada.
    """
    from .models import LegacyManual, LegacyManualRole

    # Obtener el manual más reciente de la entidad
    manual_qs = LegacyManual.objects.filter(entity=entity).order_by('-issue_date', '-created_at')
    if not manual_qs.exists():
        return {
            'added': [], 'removed': [], 'modified': [], 'unchanged': [],
            'stats': {'added': 0, 'removed': 0, 'modified': 0, 'unchanged': 0},
            'warnings': ['No hay manuales vigentes registrados para esta entidad.'],
        }

    manual = manual_qs.first()
    legacy_roles = list(LegacyManualRole.objects.filter(manual=manual))

    # Obtener el plan propuesto más reciente de la restructuring
    from apps.planta.models import PayrollPlan, PayrollPosition
    plan_qs = PayrollPlan.objects.filter(
        restructuring=restructuring,
        kind=PayrollPlan.Kind.PROPOSED,
    ).order_by('-reference_date', '-created_at')

    if not plan_qs.exists():
        return {
            'added': [], 'removed': [], 'modified': [], 'unchanged': [],
            'stats': {'added': 0, 'removed': 0, 'modified': 0, 'unchanged': 0},
            'warnings': ['No hay plan propuesto en la reestructuración.'],
        }

    plan = plan_qs.first()
    proposed_positions = list(PayrollPosition.objects.filter(plan=plan))

    def key_for_role(code: str, grade: str, denomination: str) -> tuple:
        if code and grade:
            return ('cg', code.strip(), grade.strip())
        return ('dn', _normalize(denomination))

    # Build lookup dicts
    legacy_dict: dict[tuple, LegacyManualRole] = {}
    for r in legacy_roles:
        k = key_for_role(r.code, r.grade, r.denomination)
        legacy_dict[k] = r

    proposed_dict: dict[tuple, PayrollPosition] = {}
    for p in proposed_positions:
        k = key_for_role(p.code, p.grade, p.denomination)
        proposed_dict[k] = p

    added = []
    removed = []
    modified = []
    unchanged = []

    # Positions in proposed not in legacy → added
    for k, pos in proposed_dict.items():
        if k not in legacy_dict:
            added.append({
                'code': pos.code, 'grade': pos.grade, 'denomination': pos.denomination,
            })
        else:
            # Compare denomination
            role = legacy_dict[k]
            diffs = {}
            if _normalize(role.denomination) != _normalize(pos.denomination):
                diffs['denomination'] = {'old': role.denomination, 'new': pos.denomination}
            if diffs:
                modified.append({
                    'code': pos.code, 'grade': pos.grade,
                    'denomination': pos.denomination, 'diff': diffs,
                })
            else:
                unchanged.append({'code': pos.code, 'grade': pos.grade})

    # Roles in legacy not in proposed → removed
    for k, role in legacy_dict.items():
        if k not in proposed_dict:
            removed.append({
                'code': role.code, 'grade': role.grade, 'denomination': role.denomination,
            })

    stats = {
        'added': len(added),
        'removed': len(removed),
        'modified': len(modified),
        'unchanged': len(unchanged),
    }
    return {
        'added': added,
        'removed': removed,
        'modified': modified,
        'unchanged': unchanged,
        'stats': stats,
    }
