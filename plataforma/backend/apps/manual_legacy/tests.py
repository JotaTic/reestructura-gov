"""Tests para apps.manual_legacy."""
from __future__ import annotations

import io

from django.contrib.auth.models import User
from django.test import TestCase

from apps.core.models import Entity, Restructuring
from apps.planta.models import PayrollPlan, PayrollPosition
from .models import LegacyManual, LegacyManualRole, LegacyManualFunction
from .services import parse_manual_docx, compare_current_vs_proposed


def _make_entity(name='TestEntity'):
    return Entity.objects.create(
        name=name, acronym=name[:4].upper(),
        order='MUNICIPAL', legal_nature='ALCALDIA',
        municipality_category='6',
    )


def _make_manual_docx_in_memory(num_roles: int = 3) -> io.BytesIO:
    """Genera un DOCX en memoria con `num_roles` cargos ficticios."""
    from docx import Document
    doc = Document()
    levels = ['PROFESIONAL', 'TECNICO', 'ASISTENCIAL']
    codes = ['219', '314', '407']
    grades = ['01', '05', '10']
    denoms = ['Profesional Universitario', 'Técnico Operativo', 'Auxiliar Administrativo']

    for i in range(num_roles):
        doc.add_heading('IDENTIFICACION DEL EMPLEO', level=2)
        doc.add_paragraph(f'NIVEL: {levels[i % 3]}')
        doc.add_paragraph(f'CODIGO: {codes[i % 3]}')
        doc.add_paragraph(f'GRADO: {grades[i % 3]}')
        doc.add_paragraph(f'DENOMINACION DEL EMPLEO: {denoms[i % 3]}')

        doc.add_heading('PROPOSITO PRINCIPAL', level=2)
        doc.add_paragraph(f'Ejecutar actividades relacionadas con {denoms[i % 3].lower()}.')

        doc.add_heading('FUNCIONES ESENCIALES', level=2)
        doc.add_paragraph(f'1. Primera función del cargo {i + 1}.')
        doc.add_paragraph(f'2. Segunda función del cargo {i + 1}.')

        doc.add_heading('REQUISITOS', level=2)
        doc.add_paragraph('ESTUDIOS: Título universitario en área relacionada.')
        doc.add_paragraph('EXPERIENCIA: Un año de experiencia profesional.')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


class ParseManualDocxTest(TestCase):
    def setUp(self):
        self.entity = _make_entity()
        self.manual = LegacyManual.objects.create(
            entity=self.entity, name='Manual 2018', issue_date='2018-01-01'
        )

    def test_creates_roles_from_docx(self):
        buf = _make_manual_docx_in_memory(3)
        result = parse_manual_docx(buf, self.manual)
        self.assertGreaterEqual(result['roles_created'], 3)
        self.assertGreaterEqual(LegacyManualRole.objects.filter(manual=self.manual).count(), 3)

    def test_creates_functions_from_docx(self):
        buf = _make_manual_docx_in_memory(3)
        result = parse_manual_docx(buf, self.manual)
        self.assertGreaterEqual(result['functions_created'], 3)
        self.assertGreaterEqual(LegacyManualFunction.objects.filter(role__manual=self.manual).count(), 3)

    def test_import_report_saved_to_manual(self):
        buf = _make_manual_docx_in_memory(1)
        parse_manual_docx(buf, self.manual)
        self.manual.refresh_from_db()
        self.assertIn('roles_created', self.manual.import_report)

    def test_result_has_expected_keys(self):
        buf = _make_manual_docx_in_memory(2)
        result = parse_manual_docx(buf, self.manual)
        self.assertIn('roles_created', result)
        self.assertIn('functions_created', result)
        self.assertIn('warnings', result)


class CompareCurrentVsProposedTest(TestCase):
    def setUp(self):
        self.entity = _make_entity('CompEntity')
        self.restructuring = Restructuring.objects.create(
            entity=self.entity,
            name='R1', code='R001',
            reference_date='2026-01-01',
            status='DRAFT',
        )
        self.manual = LegacyManual.objects.create(
            entity=self.entity, name='Manual 2020', issue_date='2020-01-01'
        )
        # Manual has 3 roles
        self.role1 = LegacyManualRole.objects.create(
            manual=self.manual, level='PROFESIONAL',
            code='219', grade='01', denomination='Profesional Universitario', order=0,
        )
        self.role2 = LegacyManualRole.objects.create(
            manual=self.manual, level='TECNICO',
            code='314', grade='05', denomination='Técnico Operativo', order=1,
        )
        self.role3 = LegacyManualRole.objects.create(
            manual=self.manual, level='ASISTENCIAL',
            code='407', grade='10', denomination='Auxiliar Administrativo', order=2,
        )

        # Proposed plan: role1 (unchanged), role2 (denomination changed), role_new (added)
        self.plan = PayrollPlan.objects.create(
            entity=self.entity, restructuring=self.restructuring,
            kind=PayrollPlan.Kind.PROPOSED,
            name='Planta Propuesta 2026',
            reference_date='2026-01-01',
        )
        # Unchanged
        PayrollPosition.objects.create(
            plan=self.plan, hierarchy_level='PROFESIONAL',
            code='219', grade='01', denomination='Profesional Universitario',
        )
        # Modified denomination
        PayrollPosition.objects.create(
            plan=self.plan, hierarchy_level='TECNICO',
            code='314', grade='05', denomination='Técnico Operativo Especializado',
        )
        # New cargo (not in manual)
        PayrollPosition.objects.create(
            plan=self.plan, hierarchy_level='DIRECTIVO',
            code='090', grade='01', denomination='Gerente',
        )
        # role3 (407-10) NOT in proposed → removed

    def test_added_contains_new_cargo(self):
        result = compare_current_vs_proposed(self.entity, self.restructuring)
        added_codes = [a['code'] for a in result['added']]
        self.assertIn('090', added_codes)

    def test_removed_contains_missing_cargo(self):
        result = compare_current_vs_proposed(self.entity, self.restructuring)
        removed_codes = [r['code'] for r in result['removed']]
        self.assertIn('407', removed_codes)

    def test_modified_contains_changed_denomination(self):
        result = compare_current_vs_proposed(self.entity, self.restructuring)
        modified_codes = [m['code'] for m in result['modified']]
        self.assertIn('314', modified_codes)

    def test_unchanged_contains_identical_cargo(self):
        result = compare_current_vs_proposed(self.entity, self.restructuring)
        unchanged_codes = [u['code'] for u in result['unchanged']]
        self.assertIn('219', unchanged_codes)

    def test_stats_structure(self):
        result = compare_current_vs_proposed(self.entity, self.restructuring)
        for key in ('added', 'removed', 'modified', 'unchanged'):
            self.assertIn(key, result['stats'])

    def test_no_manual_returns_warnings(self):
        entity2 = _make_entity('E2')
        r2 = Restructuring.objects.create(
            entity=entity2, name='R2', code='R002',
            reference_date='2026-01-01', status='DRAFT',
        )
        result = compare_current_vs_proposed(entity2, r2)
        self.assertIn('warnings', result)
