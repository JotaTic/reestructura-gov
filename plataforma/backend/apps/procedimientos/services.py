"""
Servicios para apps.procedimientos.

- parse_procedure_docx: importa pasos desde DOCX.
- derive_workload_from_procedures: propone entradas de cargas desde procedimientos.
"""
from __future__ import annotations

import re
import unicodedata


def _normalize(text: str) -> str:
    return unicodedata.normalize('NFKD', text).encode('ascii', 'ignore').decode().lower().strip()


def _parse_minutes(text: str) -> int:
    """Convierte '15 min', '1 hora', '1h', '30' a minutos enteros."""
    t = text.lower().strip()
    if 'hora' in t or t.endswith('h'):
        m = re.search(r'(\d+(?:\.\d+)?)', t)
        if m:
            return int(float(m.group(1)) * 60)
    m = re.search(r'(\d+)', t)
    return int(m.group(1)) if m else 0


# ---------------------------------------------------------------------------
# parse_procedure_docx
# ---------------------------------------------------------------------------

def parse_procedure_docx(file, process_id: int) -> dict:
    """
    Importa un procedimiento desde un DOCX.

    Estrategia:
    1. Título del procedimiento = primer párrafo de encabezado o primer párrafo no vacío.
    2. Metadatos: busca párrafos/tablas con "objetivo", "alcance", "entradas", "salidas".
    3. Primera tabla con ≥3 columnas con palabras clave de pasos → ProcedureStep por fila.
    4. Minutos: parsea strings "15 min", "1 hora", etc.
    """
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from .models import Procedure, ProcedureStep
    from apps.procesos.models import Process

    doc = DocxDocument(file)
    warnings: list[str] = []
    steps_created = 0

    try:
        process = Process.objects.get(pk=process_id)
    except Process.DoesNotExist:
        return {'procedure_id': None, 'steps_created': 0,
                'warnings': [f'Proceso con id {process_id} no encontrado.']}

    # --- Extract title from first non-empty paragraph ---
    title = ''
    for para in doc.paragraphs:
        if para.text.strip():
            title = para.text.strip()
            break
    if not title:
        title = 'Procedimiento importado'

    # --- Extract metadata from paragraphs ---
    objective = ''
    scope = ''
    inputs_text = ''
    outputs_text = ''

    KW_MAP = {
        'objetivo': 'objective',
        'alcance': 'scope',
        'entradas': 'inputs',
        'salidas': 'outputs',
    }
    current_key = None
    buffer = []

    def flush_meta():
        nonlocal objective, scope, inputs_text, outputs_text, buffer, current_key
        text_val = '\n'.join(buffer).strip()
        if current_key == 'objective':
            objective = text_val
        elif current_key == 'scope':
            scope = text_val
        elif current_key == 'inputs':
            inputs_text = text_val
        elif current_key == 'outputs':
            outputs_text = text_val
        buffer.clear()
        current_key = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        norm = _normalize(text)
        matched = None
        for kw, fld in KW_MAP.items():
            if kw in norm:
                matched = fld
                break
        if matched:
            flush_meta()
            current_key = matched
            # Value might be on same line after ':'
            if ':' in text:
                rest = text.split(':', 1)[1].strip()
                if rest:
                    buffer.append(rest)
        elif current_key:
            buffer.append(text)

    flush_meta()

    # --- Find step table ---
    STEP_KEYWORDS = {'no', 'paso', 'actividad', 'responsable', 'tiempo', 'descripcion',
                     'cargo', 'minutos', 'entrada', 'salida', 'sistema'}

    step_table = None
    col_map: dict[str, int] = {}

    for table in doc.tables:
        if len(table.columns) < 3:
            continue
        # Check header row
        headers = [_normalize(c.text) for c in table.rows[0].cells]
        score = sum(1 for h in headers if any(kw in h for kw in STEP_KEYWORDS))
        if score >= 2:
            step_table = table
            for i, h in enumerate(headers):
                if any(kw in h for kw in ('no', 'paso', '#', 'item')):
                    col_map.setdefault('order', i)
                elif any(kw in h for kw in ('actividad', 'descripcion', 'paso')):
                    col_map.setdefault('description', i)
                elif any(kw in h for kw in ('responsable', 'cargo', 'ejecutor')):
                    col_map.setdefault('role_executor', i)
                elif any(kw in h for kw in ('tiempo', 'minutos', 'duracion')):
                    col_map.setdefault('estimated_minutes', i)
                elif 'entrada' in h:
                    col_map.setdefault('input_document', i)
                elif 'salida' in h:
                    col_map.setdefault('output_document', i)
                elif 'sistema' in h:
                    col_map.setdefault('supporting_system', i)
            break

    # Create Procedure
    procedure = Procedure.objects.create(
        process=process,
        code=f'PROC-{process.id}-IMP',
        name=title,
        version='1.0',
        objective=objective,
        scope=scope,
        inputs_text=inputs_text,
        outputs_text=outputs_text,
    )

    # Create steps from table
    if step_table is not None:
        desc_col = col_map.get('description', 1)
        for row_idx, row in enumerate(step_table.rows[1:], start=1):
            cells = row.cells
            n_cells = len(cells)
            description = cells[desc_col].text.strip() if desc_col < n_cells else ''
            if not description:
                continue
            role_executor = cells[col_map['role_executor']].text.strip() if 'role_executor' in col_map and col_map['role_executor'] < n_cells else ''
            time_text = cells[col_map['estimated_minutes']].text.strip() if 'estimated_minutes' in col_map and col_map['estimated_minutes'] < n_cells else '0'
            est_min = _parse_minutes(time_text)
            input_doc = cells[col_map['input_document']].text.strip() if 'input_document' in col_map and col_map['input_document'] < n_cells else ''
            output_doc = cells[col_map['output_document']].text.strip() if 'output_document' in col_map and col_map['output_document'] < n_cells else ''
            system = cells[col_map['supporting_system']].text.strip() if 'supporting_system' in col_map and col_map['supporting_system'] < n_cells else ''

            ProcedureStep.objects.create(
                procedure=procedure,
                order=row_idx,
                description=description,
                role_executor=role_executor,
                estimated_minutes=est_min,
                input_document=input_doc,
                output_document=output_doc,
                supporting_system=system,
            )
            steps_created += 1
    else:
        warnings.append('No se encontró tabla de pasos con columnas reconocibles.')

    return {
        'procedure_id': procedure.pk,
        'steps_created': steps_created,
        'warnings': warnings,
    }


# ---------------------------------------------------------------------------
# derive_workload_from_procedures
# ---------------------------------------------------------------------------

def derive_workload_from_procedures(process_map) -> list[dict]:
    """
    Propone entradas de matriz de cargas basadas en pasos de procedimientos.

    Para cada proceso del mapa que tenga procedimientos, por cada paso con
    estimated_minutes > 0, genera un dict compatible con WorkloadEntry.

    NO persiste. El usuario confirma en el frontend antes de guardar.
    """
    from apps.procesos.models import Process

    proposals = []
    processes = Process.objects.filter(process_map=process_map).prefetch_related(
        'procedures__steps'
    )

    for process in processes:
        for procedure in process.procedures.all():
            for step in procedure.steps.all():
                if step.estimated_minutes <= 0:
                    continue
                # Convert estimated_minutes to hours
                est_hours = round(step.estimated_minutes / 60.0, 4)
                monthly_vol = step.monthly_volume or 1
                proposals.append({
                    'process': process.name,
                    'activity': step.description[:200],
                    'procedure': procedure.name,
                    'procedure_step_id': step.pk,
                    'department_id': None,  # usuario debe asignar
                    'job_denomination': step.role_executor or '',
                    'job_code': '',
                    'job_grade': '',
                    'monthly_frequency': monthly_vol,
                    't_min': est_hours,
                    't_usual': est_hours,
                    't_max': round(est_hours * 1.2, 4),
                })

    return proposals
