"""Tests para apps.procedimientos."""
from __future__ import annotations

import io

from django.test import TestCase

from apps.core.models import Entity, Restructuring
from apps.procesos.models import ProcessMap, Process
from .models import Procedure, ProcedureStep
from .services import parse_procedure_docx, derive_workload_from_procedures


def _make_entity(name='ProcEntity'):
    return Entity.objects.create(
        name=name, acronym=name[:4].upper(),
        order='MUNICIPAL', legal_nature='ALCALDIA',
        municipality_category='6',
    )


def _make_process(entity, restructuring, name='Gestión Documental'):
    pm = ProcessMap.objects.create(
        entity=entity, restructuring=restructuring,
        kind=ProcessMap.Kind.CURRENT, name='Mapa de Procesos',
        reference_date='2026-01-01',
    )
    process = Process.objects.create(
        process_map=pm, code='P01', name=name,
        type=Process.Type.MISIONAL, order=0,
    )
    return pm, process


def _make_procedure_docx_in_memory(num_steps: int = 6) -> io.BytesIO:
    """Genera un DOCX con tabla de pasos."""
    from docx import Document
    doc = Document()
    doc.add_heading('Procedimiento de Gestión de Correspondencia', level=1)
    doc.add_paragraph('Objetivo: Gestionar la correspondencia de la entidad.')
    doc.add_paragraph('Alcance: Aplica a todas las dependencias.')
    doc.add_paragraph('Entradas: Documentos recibidos.')
    doc.add_paragraph('Salidas: Documentos distribuidos y radicados.')

    # Add steps table
    table = doc.add_table(rows=1 + num_steps, cols=5)
    headers = ['No.', 'Actividad', 'Responsable', 'Tiempo', 'Sistema']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    roles = ['Auxiliar Administrativo', 'Profesional', 'Técnico', 'Director', 'Asesor', 'Auxiliar']
    for i in range(num_steps):
        row = table.rows[i + 1]
        row.cells[0].text = str(i + 1)
        row.cells[1].text = f'Actividad {i + 1}: Ejecutar tarea correspondiente al paso {i + 1}.'
        row.cells[2].text = roles[i % len(roles)]
        row.cells[3].text = f'{15 + i * 5} min'
        row.cells[4].text = 'SGDEA'

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


class ParseProcedureDocxTest(TestCase):
    def setUp(self):
        self.entity = _make_entity()
        self.restructuring = Restructuring.objects.create(
            entity=self.entity, name='R1', code='R001',
            reference_date='2026-01-01', status='DRAFT',
        )
        self.pm, self.process = _make_process(self.entity, self.restructuring)

    def test_creates_procedure_and_steps(self):
        buf = _make_procedure_docx_in_memory(6)
        result = parse_procedure_docx(buf, self.process.pk)
        self.assertIsNotNone(result['procedure_id'])
        self.assertEqual(result['steps_created'], 6)

    def test_procedure_linked_to_process(self):
        buf = _make_procedure_docx_in_memory(6)
        result = parse_procedure_docx(buf, self.process.pk)
        proc = Procedure.objects.get(pk=result['procedure_id'])
        self.assertEqual(proc.process_id, self.process.pk)

    def test_steps_have_estimated_minutes(self):
        buf = _make_procedure_docx_in_memory(6)
        result = parse_procedure_docx(buf, self.process.pk)
        steps = ProcedureStep.objects.filter(procedure_id=result['procedure_id'])
        for step in steps:
            self.assertGreater(step.estimated_minutes, 0)

    def test_invalid_process_returns_warning(self):
        buf = _make_procedure_docx_in_memory(1)
        result = parse_procedure_docx(buf, 99999)
        self.assertIsNone(result['procedure_id'])
        self.assertTrue(len(result['warnings']) > 0)


class DeriveWorkloadFromProceduresTest(TestCase):
    def setUp(self):
        self.entity = _make_entity('DWEntity')
        self.restructuring = Restructuring.objects.create(
            entity=self.entity, name='R2', code='R002',
            reference_date='2026-01-01', status='DRAFT',
        )
        self.pm, self.process = _make_process(self.entity, self.restructuring, 'Proceso A')

        # Create procedure with steps
        self.procedure = Procedure.objects.create(
            process=self.process, code='P-001', name='Procedimiento A', version='1.0',
        )
        for i in range(3):
            ProcedureStep.objects.create(
                procedure=self.procedure, order=i,
                description=f'Paso {i + 1}',
                estimated_minutes=30 + i * 10,
                monthly_volume=20,
            )

    def test_returns_proposals_proportional_to_steps(self):
        proposals = derive_workload_from_procedures(self.pm)
        self.assertEqual(len(proposals), 3)

    def test_proposals_have_procedure_step_id(self):
        proposals = derive_workload_from_procedures(self.pm)
        for p in proposals:
            self.assertIsNotNone(p.get('procedure_step_id'))

    def test_proposals_have_positive_minutes(self):
        proposals = derive_workload_from_procedures(self.pm)
        for p in proposals:
            self.assertGreater(p['t_usual'], 0)

    def test_step_without_minutes_excluded(self):
        """Pasos con estimated_minutes=0 no se proponen."""
        self.procedure.steps.all().update(estimated_minutes=0)
        proposals = derive_workload_from_procedures(self.pm)
        self.assertEqual(len(proposals), 0)
