"""
Helpers genéricos de exportación a XLSX y DOCX para los módulos de
ReEstructura.Gov.

Modelo de datos:
    Un export = (title, meta, sections)
      - title  : str           -> Título principal del documento.
      - meta   : list[tuple]   -> Pares clave/valor que se imprimen en el
                                   encabezado del documento (Entidad, Fecha,
                                   Reestructuración, etc.).
      - sections: list[Section]-> Cada sección es una "hoja" en Excel y una
                                   tabla en Word. Puede incluir una descripción
                                   libre sobre la tabla.

    Section es un TypedDict con:
      - heading     : str               — nombre de la hoja / título en Word
      - description : str (opcional)    — párrafo descriptivo
      - headers     : list[str]         — cabecera de tabla
      - rows        : list[list[Any]]   — filas de datos
      - notes       : str (opcional)    — párrafo final

Los módulos individuales solo deben producir esa estructura desde sus datos;
este helper se encarga del estilo, formato y respuesta HTTP.
"""
from __future__ import annotations

from io import BytesIO
from typing import Any, TypedDict

from django.http import HttpResponse
from rest_framework.renderers import BaseRenderer, JSONRenderer


XLSX_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
DOCX_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'


class BinaryFileRenderer(BaseRenderer):
    """
    Renderer "pasa-todo" para los @action que devuelven un HttpResponse con
    un archivo binario (xlsx/docx) directamente.

    Por qué existe: DRF ejecuta `perform_content_negotiation` ANTES de entrar
    al código del @action, y rechaza con 406 cualquier petición cuyo header
    `Accept` no sea servible por algún renderer registrado. El navegador, al
    descargar un archivo, manda `Accept: application/vnd.openxmlformats-...`;
    sin este renderer, DRF responde 406 sin ejecutar la vista.

    Con `media_type = '*/*'`, la content negotiation queda satisfecha para
    cualquier Accept, y como la vista devuelve un `HttpResponse` directo
    (no un `rest_framework.response.Response`), el método `render` de este
    renderer en realidad nunca se invoca — solo sirve para pasar la
    negociación.
    """
    media_type = '*/*'
    format = 'file'
    charset = None
    render_style = 'binary'

    def render(self, data, accepted_media_type=None, renderer_context=None):
        if isinstance(data, (bytes, bytearray)):
            return bytes(data)
        return b''


# Tupla lista para usar en `@action(..., renderer_classes=EXPORT_RENDERERS)`.
# Incluye JSONRenderer primero para que los errores (ValidationError, 403,
# 404) se sigan devolviendo como JSON cuando el cliente pida explícitamente
# `Accept: application/json`.
EXPORT_RENDERERS = (JSONRenderer, BinaryFileRenderer)

BRAND_COLOR = '0E7490'  # ReEstructura.Gov cyan-700


class Section(TypedDict, total=False):
    heading: str
    description: str
    headers: list[str]
    rows: list[list[Any]]
    notes: str


# ---------------------------------------------------------------------------
# Utilidades
# ---------------------------------------------------------------------------

def _safe_sheet_name(name: str, used: set[str]) -> str:
    """Excel limita a 31 chars y no permite ciertos caracteres."""
    banned = '[]:*?/\\'
    clean = ''.join(c for c in (name or 'Hoja') if c not in banned).strip() or 'Hoja'
    clean = clean[:31]
    base = clean
    i = 2
    while clean in used:
        suffix = f' ({i})'
        clean = (base[: 31 - len(suffix)] + suffix)
        i += 1
    used.add(clean)
    return clean


def _stringify(value: Any) -> str:
    if value is None:
        return ''
    if isinstance(value, (str, int, float)):
        return str(value)
    return str(value)


# ---------------------------------------------------------------------------
# Excel (.xlsx)
# ---------------------------------------------------------------------------

def build_xlsx(title: str, meta: list[tuple[str, str]], sections: list[Section]) -> bytes:
    """
    Construye un .xlsx con:
      - Una primera hoja "Resumen" con título y metadata.
      - Una hoja por sección; si una sección no tiene filas, escribe
        "Sin datos registrados".
    Devuelve los bytes del archivo.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)

    title_font = Font(bold=True, size=14, color='FFFFFF')
    title_fill = PatternFill('solid', fgColor=BRAND_COLOR)
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill('solid', fgColor=BRAND_COLOR)
    meta_label_font = Font(bold=True)
    wrap_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    wrap_left = Alignment(horizontal='left', vertical='top', wrap_text=True)

    used_names: set[str] = set()

    # ---- Resumen ----
    ws = wb.create_sheet(title=_safe_sheet_name('Resumen', used_names))
    ws.merge_cells('A1:D1')
    ws['A1'] = title or 'Reporte'
    ws['A1'].font = title_font
    ws['A1'].fill = title_fill
    ws['A1'].alignment = wrap_center
    ws.row_dimensions[1].height = 26

    row_idx = 3
    for label, value in (meta or []):
        ws.cell(row=row_idx, column=1, value=label).font = meta_label_font
        ws.cell(row=row_idx, column=2, value=_stringify(value)).alignment = wrap_left
        row_idx += 1
    for col_letter, width in zip('ABCD', [28, 60, 20, 20]):
        ws.column_dimensions[col_letter].width = width

    # ---- Secciones ----
    for section in sections or []:
        heading = section.get('heading') or 'Sección'
        headers = section.get('headers') or []
        rows = section.get('rows') or []
        description = section.get('description')
        notes = section.get('notes')

        ws = wb.create_sheet(title=_safe_sheet_name(heading, used_names))
        r = 1
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=max(len(headers) or 1, 1))
        ws.cell(row=r, column=1, value=heading).font = title_font
        ws.cell(row=r, column=1).fill = title_fill
        ws.cell(row=r, column=1).alignment = wrap_center
        ws.row_dimensions[r].height = 22
        r += 1

        if description:
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=max(len(headers) or 1, 1))
            ws.cell(row=r, column=1, value=description).alignment = wrap_left
            r += 1

        r += 1  # blank

        if not headers:
            ws.cell(row=r, column=1, value='(Sin columnas definidas)').font = Font(italic=True, color='888888')
            continue

        # cabecera
        for c, h in enumerate(headers, start=1):
            cell = ws.cell(row=r, column=c, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = wrap_center
        ws.row_dimensions[r].height = 28
        header_row = r
        r += 1

        # filas
        if not rows:
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=len(headers))
            ws.cell(row=r, column=1, value='Sin datos registrados').font = Font(italic=True, color='888888')
            r += 1
        else:
            for row_data in rows:
                for c, v in enumerate(row_data, start=1):
                    cell = ws.cell(row=r, column=c, value=v if isinstance(v, (int, float)) else _stringify(v))
                    cell.alignment = wrap_left
                r += 1

        if notes:
            r += 1
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=len(headers))
            ws.cell(row=r, column=1, value=notes).font = Font(italic=True, color='555555')

        # Ajuste de anchos (heurística por contenido de la cabecera y filas)
        for c_idx in range(1, len(headers) + 1):
            col_letter = get_column_letter(c_idx)
            max_len = len(_stringify(headers[c_idx - 1]))
            for row_data in rows[:200]:  # sampling
                if c_idx - 1 < len(row_data):
                    max_len = max(max_len, len(_stringify(row_data[c_idx - 1])))
            ws.column_dimensions[col_letter].width = min(max(max_len + 2, 12), 60)
        # freeze header
        ws.freeze_panes = ws.cell(row=header_row + 1, column=1)

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Word (.docx)
# ---------------------------------------------------------------------------

def build_docx(title: str, meta: list[tuple[str, str]], sections: list[Section]) -> bytes:
    """
    Construye un .docx con:
      - Portada (título + metadata en tabla de 2 columnas).
      - Una sección por cada `Section`, con heading 1, descripción opcional,
        tabla de datos y notas opcionales.
    """
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Título
    h = doc.add_heading(title or 'Reporte', level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0E, 0x74, 0x90)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tabla de metadata
    if meta:
        mt = doc.add_table(rows=len(meta), cols=2)
        mt.autofit = True
        for i, (label, value) in enumerate(meta):
            cells = mt.rows[i].cells
            cells[0].text = str(label or '')
            cells[1].text = _stringify(value)
            for run in cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
            for run in cells[1].paragraphs[0].runs:
                run.font.size = Pt(10)
        doc.add_paragraph()

    def _set_cell_bg(cell, color_hex: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tc_pr.append(shd)

    for section in sections or []:
        heading = section.get('heading') or 'Sección'
        headers = section.get('headers') or []
        rows = section.get('rows') or []
        description = section.get('description')
        notes = section.get('notes')

        sh = doc.add_heading(heading, level=1)
        for run in sh.runs:
            run.font.color.rgb = RGBColor(0x0E, 0x74, 0x90)

        if description:
            p = doc.add_paragraph(description)
            for run in p.runs:
                run.font.size = Pt(10)

        if not headers:
            doc.add_paragraph('(Sin columnas definidas)').italic = True
            continue

        table = doc.add_table(rows=1 + max(len(rows), 1), cols=len(headers))
        table.style = 'Light List Accent 1'

        # cabecera
        header_cells = table.rows[0].cells
        for c, h_text in enumerate(headers):
            header_cells[c].text = str(h_text)
            _set_cell_bg(header_cells[c], BRAND_COLOR)
            for p in header_cells[c].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            header_cells[c].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # cuerpo
        if not rows:
            only = table.rows[1].cells
            # merge horizontal usando primer cell y borrando texto de los demás
            merged = only[0]
            for c in only[1:]:
                merged = merged.merge(c)
            merged.text = 'Sin datos registrados'
            for p in merged.paragraphs:
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                    run.font.size = Pt(10)
        else:
            for i, row_data in enumerate(rows, start=1):
                cells = table.rows[i].cells
                for c, v in enumerate(row_data):
                    if c < len(cells):
                        cells[c].text = _stringify(v)
                        for p in cells[c].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)

        if notes:
            p = doc.add_paragraph(notes)
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# HttpResponse helpers
# ---------------------------------------------------------------------------

def _safe_filename_part(text: str) -> str:
    import re
    text = (text or '').strip().lower()
    text = re.sub(r'[^a-z0-9]+', '_', text)
    return text.strip('_') or 'reporte'


def xlsx_response(content: bytes, base_name: str, *context: str) -> HttpResponse:
    parts = [_safe_filename_part(base_name)] + [_safe_filename_part(c) for c in context if c]
    filename = '_'.join(p for p in parts if p) + '.xlsx'
    resp = HttpResponse(content, content_type=XLSX_CONTENT_TYPE)
    resp['Content-Disposition'] = f'attachment; filename="{filename}"'
    return resp


def docx_response(content: bytes, base_name: str, *context: str) -> HttpResponse:
    parts = [_safe_filename_part(base_name)] + [_safe_filename_part(c) for c in context if c]
    filename = '_'.join(p for p in parts if p) + '.docx'
    resp = HttpResponse(content, content_type=DOCX_CONTENT_TYPE)
    resp['Content-Disposition'] = f'attachment; filename="{filename}"'
    return resp


# ---------------------------------------------------------------------------
# Dispatcher: construye el export en el formato pedido
# ---------------------------------------------------------------------------

def export_response(
    fmt: str,
    title: str,
    meta: list[tuple[str, str]],
    sections: list[Section],
    filename_base: str,
    *filename_context: str,
) -> HttpResponse:
    """
    Fachada: recibe el formato pedido ('xlsx' | 'docx') y devuelve la
    HttpResponse lista para retornar desde un @action del ViewSet.
    """
    fmt = (fmt or '').lower()
    if fmt == 'xlsx':
        content = build_xlsx(title, meta, sections)
        return xlsx_response(content, filename_base, *filename_context)
    if fmt == 'docx':
        content = build_docx(title, meta, sections)
        return docx_response(content, filename_base, *filename_context)
    from rest_framework.exceptions import ValidationError
    raise ValidationError({'format': f"Formato no soportado: {fmt!r}. Usa 'xlsx' o 'docx'."})
